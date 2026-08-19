"""Import de diagnostics existants au format Word (.docx), pour des documents à
structure libre (notes de terrain, trames différentes de l'étoile du conseil).

Contrairement à l'import Excel/CSV (modules/import_data.py, qui exige le modèle
exact), ce module accepte un ou plusieurs fichiers Word à la structure quelconque
et utilise l'IA pour en extraire les informations vers le schéma de l'étoile du
conseil. Comme pour le chat conversationnel (agent/orchestrator.py) :

- Nécessite une connexion Internet ET une clé API (ANTHROPIC_API_KEY) — sinon
  l'appelant doit proposer le repli vers l'import Excel/CSV ou la saisie manuelle.
- Ne remplit jamais un champ sans base dans le texte fourni.
- Le résultat n'est JAMAIS sauvegardé automatiquement : il doit être relu et
  corrigé par le conseiller dans le formulaire avant tout enregistrement.
"""
import json
import os

from docx import Document as DocxDocument

from modules.collecte import load_schema


def word_import_available() -> bool:
    return bool(os.environ.get("ANTHROPIC_API_KEY"))


def extract_text_from_docx(file) -> str:
    """Extrait tout le texte utile (paragraphes + tableaux) d'un fichier .docx.
    `file` est un objet file-like (ex. UploadedFile de Streamlit)."""
    doc = DocxDocument(file)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells_text = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells_text:
                parts.append(" | ".join(cells_text))

    return "\n".join(parts)


def _build_extraction_tool() -> dict:
    schema = load_schema()["branches"]
    branch_properties = {}
    for branch_key, branch in schema.items():
        field_props = {}
        for field in branch["fields"]:
            if field["type"] == "activity_list":
                field_props[field["id"]] = {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "nom": {"type": "string"},
                            "part_marche_relative": {"type": "number"},
                            "taux_croissance": {"type": "number"},
                        },
                    },
                }
            elif field["type"] == "number":
                field_props[field["id"]] = {"type": "number"}
            else:
                field_props[field["id"]] = {"type": "string"}
        branch_properties[branch_key] = {"type": "object", "properties": field_props}

    return {
        "name": "set_diagnostic",
        "description": (
            "Enregistre les informations extraites du document dans la structure "
            "de l'étoile du conseil et, si le document en contient, dans une "
            "analyse SWOT explicite. N'invente et ne déduis jamais une information "
            "absente du texte fourni : laisse un champ vide plutôt que de deviner."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "etoile": {"type": "object", "properties": branch_properties},
                "swot": {
                    "type": "object",
                    "properties": {
                        "forces": {"type": "array", "items": {"type": "string"}},
                        "faiblesses": {"type": "array", "items": {"type": "string"}},
                        "opportunites": {"type": "array", "items": {"type": "string"}},
                        "menaces": {"type": "array", "items": {"type": "string"}},
                    },
                    "description": "Uniquement si le document contient explicitement "
                                    "des sections Forces/Faiblesses/Opportunités/Menaces "
                                    "ou équivalent (atouts/contraintes...).",
                },
            },
            "required": ["etoile"],
        },
    }


def extract_diagnostic_from_text(text: str, lang: str = "fr") -> dict:
    """Envoie le texte extrait d'un ou plusieurs documents Word à l'API Anthropic
    et retourne {"etoile": {...}, "swot": {...}}. Lève une exception si l'appel
    échoue (l'appelant doit proposer un repli vers la saisie manuelle)."""
    import anthropic  # import local : dépendance optionnelle, seulement utile en ligne

    client = anthropic.Anthropic()
    tool = _build_extraction_tool()
    schema_summary = {
        key: [f["id"] for f in branch["fields"]]
        for key, branch in load_schema()["branches"].items()
    }

    system_prompt = (
        "Tu aides un conseiller agropastoral à convertir un diagnostic existant "
        "(rédigé librement, en Word) vers une structure standardisée : l'étoile "
        "du conseil (6 branches, champs disponibles ci-après) et, si présent, une "
        f"analyse SWOT explicite. Champs disponibles par branche : {json.dumps(schema_summary, ensure_ascii=False)}. "
        "Le document peut être en plusieurs parties (ex. diagnostic général + "
        "diagnostic économique et financier) : synthétise l'ensemble en une seule "
        "structure cohérente. N'invente rien : si une information n'apparaît pas "
        "clairement dans le texte, laisse le champ correspondant vide. "
        "Utilise l'outil set_diagnostic pour retourner le résultat."
    )

    # Les documents de terrain peuvent être longs : on tronque prudemment pour
    # rester dans une fenêtre de contexte raisonnable tout en gardant l'essentiel.
    truncated_text = text[:60000]

    response = client.messages.create(
        model="claude-sonnet-5",
        max_tokens=4000,
        system=system_prompt,
        tools=[tool],
        tool_choice={"type": "tool", "name": "set_diagnostic"},
        messages=[{"role": "user", "content": truncated_text}],
    )

    for block in response.content:
        if getattr(block, "type", "") == "tool_use" and block.name == "set_diagnostic":
            return block.input

    raise RuntimeError("L'extraction n'a renvoyé aucun résultat exploitable.")


def _sanitize_etoile(raw_etoile) -> dict:
    """Garantit une structure etoile propre : un dict, dont chaque branche est
    elle-même un dict. Neutralise silencieusement toute forme inattendue
    renvoyée par l'extraction IA (liste, chaîne, None...) plutôt que de
    laisser une donnée malformée se propager et faire planter l'affichage
    plus tard."""
    if not isinstance(raw_etoile, dict):
        return {}
    clean = {}
    for branch_key, branch_value in raw_etoile.items():
        clean[branch_key] = branch_value if isinstance(branch_value, dict) else {}
    return clean


def build_diagnostic_from_extraction(extraction: dict, nom: str, type_structure: str,
                                      conseiller: str) -> dict:
    """Construit un diagnostic complet à partir du résultat d'extraction, prêt à
    être présenté au conseiller pour relecture (jamais sauvegardé automatiquement)."""
    diagnostic = {
        "nom": nom,
        "type": type_structure,
        "conseiller": conseiller,
        "etoile": _sanitize_etoile(extraction.get("etoile")),
    }
    swot = extraction.get("swot")
    if isinstance(swot, dict) and any(swot.get(k) for k in ("forces", "faiblesses", "opportunites", "menaces")):
        diagnostic["swot_import"] = swot
    return diagnostic
