"""Export du plan stratégique et du plan d'actions en PDF et Word.
Ces fonctions ne doivent être appelées par l'interface qu'après validation
explicite du conseiller (voir modules/plan_strategique.py::is_validated)."""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

from modules.plan_strategique import is_validated


class ExportNotAllowedError(Exception):
    """Levée si on tente d'exporter un plan non validé par le conseiller."""


def _check_validated(diagnostic: dict):
    if not is_validated(diagnostic):
        raise ExportNotAllowedError(
            "Le plan doit être validé par le conseiller avant export.")


def _identity_label(diagnostic: dict, lang: str, include_real_name: bool) -> str:
    """Protection des données : par défaut, seul le code d'identifiant du
    diagnostic apparaît dans les documents exportés — jamais le nom réel de
    l'EFA/OP, sauf si le conseiller a explicitement demandé de l'inclure
    (usage strictement interne)."""
    prefix = "Exploitation / OP" if lang == "fr" else "Farm / PO"
    code = diagnostic.get("code", "")
    if include_real_name:
        nom = diagnostic.get("nom", "")
        return f"{prefix}: {code} - {nom}"
    return f"{prefix}: {code}"


def _add_entreprise_section_word(doc, diagnostic: dict, lang: str):
    """Ajoute la section Entreprise (histoire, environnement, parcelles,
    calendrier, activités avec marges, diagnostic financier, immobilisations,
    bilan) au document Word. Données non identifiantes : toujours incluses,
    indépendamment du masquage du nom/contact."""
    from modules.entreprise import compute_marge_brute, compute_marge_directe, compute_diagnostic_financier

    ent = diagnostic.get("entreprise", {})
    if not ent:
        return

    doc.add_heading("Entreprise" if lang == "fr" else "Business", level=1)

    histoire = ent.get("histoire", [])
    if histoire:
        doc.add_heading("Histoire" if lang == "fr" else "History", level=2)
        for h in histoire:
            doc.add_paragraph(f"{h.get('date', '')} — {h.get('quoi', '')} ({h.get('pourquoi', '')})",
                               style="List Bullet")

    env = ent.get("environnement", {})
    if env.get("marche_clients_concurrents") or env.get("partenaires_fournisseurs_milieu"):
        doc.add_heading("Environnement" if lang == "fr" else "Environment", level=2)
        if env.get("marche_clients_concurrents"):
            doc.add_paragraph(env["marche_clients_concurrents"])
        if env.get("partenaires_fournisseurs_milieu"):
            doc.add_paragraph(env["partenaires_fournisseurs_milieu"])

    parcelles = ent.get("parcelles", [])
    if parcelles:
        doc.add_heading("Plan de localisation" if lang == "fr" else "Site plan", level=2)
        for p in parcelles:
            doc.add_paragraph(
                f"{p.get('nom', '')} — {p.get('zonage', '')} — {p.get('utilisation', '')} "
                f"— {p.get('production', '')} — {p.get('surface', 0)} ha — {p.get('statut', '')}"
                f"{'' if p.get('mise_en_valeur', True) else (' (non mise en valeur)' if lang == 'fr' else ' (unused)')}",
                style="List Bullet")

    calendrier = ent.get("calendrier", [])
    if calendrier:
        doc.add_heading("Calendrier des activités" if lang == "fr" else "Activity calendar", level=2)
        for c in calendrier:
            doc.add_paragraph(f"{c.get('activite', '')} ({c.get('type', '')}): {', '.join(c.get('mois', []))}",
                               style="List Bullet")

    activites = ent.get("activites", [])
    if activites:
        doc.add_heading("Description des activités" if lang == "fr" else "Activity description", level=2)
        for a in activites:
            doc.add_heading(a.get("nom", "") or "-", level=3)
            mb, md = compute_marge_brute(a), compute_marge_directe(a)
            doc.add_paragraph(
                f"{'Marge brute' if lang == 'fr' else 'Gross margin'}: {mb:,.0f} | "
                f"{'Marge directe' if lang == 'fr' else 'Direct margin'}: {md:,.0f}")
            if a.get("finalites_objectifs"):
                doc.add_paragraph(a["finalites_objectifs"])
            if a.get("points_forts"):
                doc.add_paragraph(f"{'Points forts' if lang == 'fr' else 'Strengths'}: {a['points_forts']}")
            if a.get("risques"):
                doc.add_paragraph(f"{'Risques' if lang == 'fr' else 'Risks'}: {a['risques']}")

    if ent.get("diagnostic_financier") or activites:
        results = compute_diagnostic_financier(diagnostic)
        doc.add_heading("Diagnostic économique et financier global" if lang == "fr"
                         else "Overall economic and financial diagnosis", level=2)
        doc.add_paragraph(f"{'Marge brute globale' if lang == 'fr' else 'Overall gross margin'}: "
                           f"{results['marge_brute_globale']:,.0f}")
        doc.add_paragraph(f"EBE: {results['ebe']:,.0f}")
        doc.add_paragraph(f"{'Marge de sécurité' if lang == 'fr' else 'Safety margin'}: "
                           f"{results['marge_securite']:,.0f}")

    immos = ent.get("immobilisations", [])
    if immos:
        doc.add_heading("Immobilisations" if lang == "fr" else "Assets", level=2)
        for im in immos:
            doc.add_paragraph(
                f"{im.get('categorie', '')} ({im.get('annee_acquisition', '')}) — "
                f"{'valeur actuelle' if lang == 'fr' else 'current value'}: {im.get('valeur_actuelle', 0):,.0f}",
                style="List Bullet")

    bilan = ent.get("bilan", {})
    if bilan.get("fin") and isinstance(bilan.get("fin"), dict) and "actif" in bilan.get("fin", {}):
        from modules.bilan import compute_totals, compute_tableau_financement
        totals_fin = compute_totals(bilan["fin"])
        doc.add_heading("Bilan (fin d'exercice)" if lang == "fr" else "Balance sheet (year end)", level=2)
        doc.add_paragraph(f"{'Total actif' if lang == 'fr' else 'Total assets'}: {totals_fin['total_actif']:,.0f} — "
                           f"{'Total passif' if lang == 'fr' else 'Total liabilities'}: {totals_fin['total_passif']:,.0f}")
        doc.add_paragraph(f"FDR: {totals_fin['fdr']:,.0f} | BFR: {totals_fin['bfr']:,.0f} | "
                           f"{'Trésorerie' if lang == 'fr' else 'Cash'}: {totals_fin['tresorerie']:,.0f}")
        if bilan.get("debut") and isinstance(bilan.get("debut"), dict) and "actif" in bilan.get("debut", {}):
            tf = compute_tableau_financement(diagnostic)
            doc.add_paragraph(
                f"{'Variation FDR' if lang == 'fr' else 'FDR change'}: {tf['delta_fdr']:,.0f} | "
                f"{'Variation BFR' if lang == 'fr' else 'BFR change'}: {tf['delta_bfr']:,.0f} | "
                f"{'Variation trésorerie' if lang == 'fr' else 'Cash change'}: {tf['delta_tresorerie']:,.0f}")


def _add_entreprise_section_pdf(pdf, diagnostic: dict, lang: str):
    """Équivalent PDF de _add_entreprise_section_word."""
    from modules.entreprise import compute_marge_brute, compute_marge_directe, compute_diagnostic_financier

    ent = diagnostic.get("entreprise", {})
    if not ent:
        return

    pdf.set_font("Helvetica", "B", 14)
    pdf.title_line("Entreprise" if lang == "fr" else "Business")
    pdf.ln(1)

    histoire = ent.get("histoire", [])
    if histoire:
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Histoire" if lang == "fr" else "History")
        pdf.set_font("Helvetica", size=10)
        for h in histoire:
            pdf.line(f"- {h.get('date', '')} - {h.get('quoi', '')} ({h.get('pourquoi', '')})", height=6)

    env = ent.get("environnement", {})
    if env.get("marche_clients_concurrents") or env.get("partenaires_fournisseurs_milieu"):
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Environnement" if lang == "fr" else "Environment")
        pdf.set_font("Helvetica", size=10)
        if env.get("marche_clients_concurrents"):
            pdf.line(env["marche_clients_concurrents"], height=6)
        if env.get("partenaires_fournisseurs_milieu"):
            pdf.line(env["partenaires_fournisseurs_milieu"], height=6)

    parcelles = ent.get("parcelles", [])
    if parcelles:
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Plan de localisation" if lang == "fr" else "Site plan")
        pdf.set_font("Helvetica", size=10)
        for p in parcelles:
            pdf.line(f"- {p.get('nom', '')} - {p.get('zonage', '')} - {p.get('utilisation', '')} "
                     f"- {p.get('surface', 0)} ha - {p.get('statut', '')}", height=6)

    calendrier = ent.get("calendrier", [])
    if calendrier:
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Calendrier des activites" if lang == "fr" else "Activity calendar")
        pdf.set_font("Helvetica", size=10)
        for c in calendrier:
            pdf.line(f"- {c.get('activite', '')} ({c.get('type', '')}): {', '.join(c.get('mois', []))}", height=6)

    activites = ent.get("activites", [])
    if activites:
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Description des activites" if lang == "fr" else "Activity description")
        pdf.set_font("Helvetica", size=10)
        for a in activites:
            mb, md = compute_marge_brute(a), compute_marge_directe(a)
            pdf.line(f"- {a.get('nom', '') or '-'} : marge brute {mb:,.0f} / marge directe {md:,.0f}"
                     if lang == "fr" else
                     f"- {a.get('nom', '') or '-'}: gross margin {mb:,.0f} / direct margin {md:,.0f}", height=6)

    if ent.get("diagnostic_financier") or activites:
        results = compute_diagnostic_financier(diagnostic)
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Diagnostic economique et financier" if lang == "fr" else "Financial diagnosis")
        pdf.set_font("Helvetica", size=10)
        pdf.line(f"- Marge brute globale: {results['marge_brute_globale']:,.0f}", height=6)
        pdf.line(f"- EBE: {results['ebe']:,.0f}", height=6)
        pdf.line(f"- Marge de securite: {results['marge_securite']:,.0f}"
                 if lang == "fr" else f"- Safety margin: {results['marge_securite']:,.0f}", height=6)

    immos = ent.get("immobilisations", [])
    if immos:
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Immobilisations" if lang == "fr" else "Assets")
        pdf.set_font("Helvetica", size=10)
        for im in immos:
            pdf.line(f"- {im.get('categorie', '')} ({im.get('annee_acquisition', '')}): "
                     f"{im.get('valeur_actuelle', 0):,.0f}", height=6)

    bilan = ent.get("bilan", {})
    if bilan.get("fin") and isinstance(bilan.get("fin"), dict) and "actif" in bilan.get("fin", {}):
        from modules.bilan import compute_totals, compute_tableau_financement
        totals_fin = compute_totals(bilan["fin"])
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Bilan (fin d'exercice)" if lang == "fr" else "Balance sheet (year end)")
        pdf.set_font("Helvetica", size=10)
        pdf.line(f"- Total actif: {totals_fin['total_actif']:,.0f} | Total passif: {totals_fin['total_passif']:,.0f}"
                 if lang == "fr" else
                 f"- Total assets: {totals_fin['total_actif']:,.0f} | "
                 f"Total liabilities: {totals_fin['total_passif']:,.0f}", height=6)
        pdf.line(f"- FDR: {totals_fin['fdr']:,.0f} | BFR: {totals_fin['bfr']:,.0f} | "
                 f"Tresorerie: {totals_fin['tresorerie']:,.0f}", height=6)
        if bilan.get("debut") and isinstance(bilan.get("debut"), dict) and "actif" in bilan.get("debut", {}):
            tf = compute_tableau_financement(diagnostic)
            pdf.line(f"- Variation FDR: {tf['delta_fdr']:,.0f} | Variation BFR: {tf['delta_bfr']:,.0f} | "
                     f"Variation tresorerie: {tf['delta_tresorerie']:,.0f}", height=6)
    pdf.ln(2)


def export_word_bytes(diagnostic: dict, plan: dict, lang: str = "fr", swot: dict = None,
                       include_real_name: bool = False) -> bytes:
    _check_validated(diagnostic)
    doc = Document()

    title = "Plan stratégique et plan d'actions" if lang == "fr" else "Strategic plan and action plan"
    doc.add_heading(title, level=0)

    doc.add_paragraph(_identity_label(diagnostic, lang, include_real_name))
    doc.add_paragraph(f"{'Conseiller' if lang == 'fr' else 'Advisor'}: {diagnostic.get('conseiller', '')}")

    validation = diagnostic.get("validation", {})
    doc.add_paragraph(
        (f"Validé par {validation.get('validated_by', '')} le {validation.get('date', '')}"
         if lang == "fr" else
         f"Validated by {validation.get('validated_by', '')} on {validation.get('date', '')}"))

    if swot:
        doc.add_heading("Analyse SWOT (FFOM)" if lang == "fr" else "SWOT analysis", level=1)
        swot_headers = [
            ("forces", "Forces" if lang == "fr" else "Strengths"),
            ("faiblesses", "Faiblesses" if lang == "fr" else "Weaknesses"),
            ("opportunites", "Opportunités" if lang == "fr" else "Opportunities"),
            ("menaces", "Menaces" if lang == "fr" else "Threats"),
        ]
        for key, title_key in swot_headers:
            doc.add_heading(title_key, level=2)
            items = swot.get(key, [])
            if items:
                for it in items:
                    doc.add_paragraph(it, style="List Bullet")
            else:
                doc.add_paragraph("—")

    doc.add_heading("Orientations stratégiques" if lang == "fr" else "Strategic orientations", level=1)
    for o in plan.get("orientations", []):
        doc.add_paragraph(o, style="List Bullet")

    doc.add_heading("Plan d'actions" if lang == "fr" else "Action plan", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Action", "Responsable", "Échéance", "Indicateur"] if lang == "fr" \
        else ["Action", "Responsible", "Deadline", "Indicator"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for item in plan.get("action_plan", []):
        row = table.add_row().cells
        row[0].text = item.get("action", "")
        row[1].text = item.get("responsable", "")
        row[2].text = item.get("echeance", "")
        row[3].text = item.get("indicateur", "")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


_UNICODE_REPLACEMENTS = {
    "\u2018": "'", "\u2019": "'",  # apostrophes typographiques ' '
    "\u201c": '"', "\u201d": '"',  # guillemets typographiques " "
    "\u2013": "-", "\u2014": "-",  # tirets demi/long – —
    "\u2026": "...",  # points de suspension …
    "\u00a0": " ",  # espace insécable
}


def _sanitize_for_pdf(text: str) -> str:
    """La police de base (Helvetica) du PDF ne supporte que le Latin-1 : un
    caractère hors de ce jeu (guillemet stylisé, tiret long, emoji, etc.) fait
    planter la génération. On translittère d'abord les cas les plus courants
    (copier-coller depuis Word notamment), puis on remplace tout ce qui
    resterait incompatible plutôt que de faire échouer tout l'export."""
    text = str(text)
    for src, dst in _UNICODE_REPLACEMENTS.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class _PDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 14)
        self.set_x(self.l_margin)
        self.cell(0, 10, _sanitize_for_pdf(self.title_text), new_x="LMARGIN", new_y="NEXT", align="C")
        self.ln(2)

    def line(self, text: str, height: float = 8):
        """multi_cell ne remet pas le curseur en marge gauche par défaut : on le force
        avant chaque ligne pour éviter une largeur disponible nulle/négative."""
        self.set_x(self.l_margin)
        self.multi_cell(0, height, _sanitize_for_pdf(text), new_x="LMARGIN", new_y="NEXT")

    def title_line(self, text: str):
        self.set_x(self.l_margin)
        self.cell(0, 8, _sanitize_for_pdf(text), new_x="LMARGIN", new_y="NEXT")


def export_pdf_bytes(diagnostic: dict, plan: dict, lang: str = "fr", swot: dict = None,
                      include_real_name: bool = False) -> bytes:
    _check_validated(diagnostic)

    title = "Plan strategique et plan d'actions" if lang == "fr" else "Strategic plan and action plan"
    pdf = _PDF()
    pdf.title_text = title
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    pdf.line(_identity_label(diagnostic, lang, include_real_name))
    pdf.line(f"{'Conseiller' if lang == 'fr' else 'Advisor'}: {diagnostic.get('conseiller', '')}")

    validation = diagnostic.get("validation", {})
    pdf.line(
        (f"Valide par {validation.get('validated_by', '')} le {validation.get('date', '')}"
         if lang == "fr" else
         f"Validated by {validation.get('validated_by', '')} on {validation.get('date', '')}"))
    pdf.ln(4)

    if swot:
        pdf.set_font("Helvetica", "B", 12)
        pdf.title_line("Analyse SWOT (FFOM)" if lang == "fr" else "SWOT analysis")
        pdf.set_font("Helvetica", size=10)
        swot_headers = [
            ("forces", "Forces" if lang == "fr" else "Strengths"),
            ("faiblesses", "Faiblesses" if lang == "fr" else "Weaknesses"),
            ("opportunites", "Opportunites" if lang == "fr" else "Opportunities"),
            ("menaces", "Menaces" if lang == "fr" else "Threats"),
        ]
        for key, title in swot_headers:
            pdf.set_font("Helvetica", "B", 11)
            pdf.title_line(title)
            pdf.set_font("Helvetica", size=10)
            items = swot.get(key, [])
            if items:
                for it in items:
                    pdf.line(f"- {it}", height=6)
            else:
                pdf.line("-", height=6)
        pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.title_line("Orientations strategiques" if lang == "fr" else "Strategic orientations")
    pdf.set_font("Helvetica", size=11)
    for o in plan.get("orientations", []):
        pdf.line(f"- {o}", height=7)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.title_line("Plan d'actions" if lang == "fr" else "Action plan")
    pdf.set_font("Helvetica", size=10)
    for item in plan.get("action_plan", []):
        text = f"- {item.get('action', '')} | {item.get('responsable', '')} | " \
               f"{item.get('echeance', '')} | {item.get('indicateur', '')}"
        pdf.line(text, height=7)

    return bytes(pdf.output())


# ---------------------------------------------------------------------------
# Export du diagnostic brut (données de l'étoile du conseil), indépendant du
# plan stratégique. Pas de contrainte de validation ici : ce n'est pas une
# recommandation, seulement les données collectées telles quelles.
# ---------------------------------------------------------------------------

def export_diagnostic_word_bytes(diagnostic: dict, lang: str = "fr",
                                  include_real_name: bool = False) -> bytes:
    """Exporte les données brutes du diagnostic (étoile du conseil) en Word."""
    from modules.collecte import load_schema

    doc = Document()
    title = "Diagnostic — Étoile du conseil" if lang == "fr" else "Diagnostic — Advisory star"
    doc.add_heading(title, level=0)

    doc.add_paragraph(_identity_label(diagnostic, lang, include_real_name))
    doc.add_paragraph(f"{'Type' if lang == 'fr' else 'Type'}: {diagnostic.get('type', '')}")
    doc.add_paragraph(f"{'Conseiller' if lang == 'fr' else 'Advisor'}: {diagnostic.get('conseiller', '')}")

    if include_real_name:
        ident = diagnostic.get("identification", {})
        if ident:
            doc.add_heading("Identification & localisation" if lang == "fr" else "Identification & location",
                             level=1)
            loc_fields = [
                ("Village" if lang == "fr" else "Village", ident.get("village", "")),
                ("Arrondissement" if lang == "fr" else "Sub-district", ident.get("arrondissement", "")),
                ("Code arrondissement" if lang == "fr" else "Sub-district code",
                 ident.get("code_arrondissement", "")),
                ("Département" if lang == "fr" else "Department", ident.get("departement", "")),
                ("Code département" if lang == "fr" else "Department code",
                 ident.get("code_departement", "")),
                ("Région" if lang == "fr" else "Region", ident.get("region", "")),
                ("Code région" if lang == "fr" else "Region code", ident.get("code_region", "")),
                ("Pays" if lang == "fr" else "Country", ident.get("pays", "")),
                ("Année" if lang == "fr" else "Year", ident.get("annee", "")),
            ]
            for label, value in loc_fields:
                if value:
                    doc.add_paragraph(f"{label}: {value}")
            gps = ident.get("gps", {})
            if gps.get("latitude") or gps.get("longitude"):
                doc.add_paragraph(f"GPS: {gps.get('latitude', '')}, {gps.get('longitude', '')}")

            from modules.identification import contact_is_visible
            contact = ident.get("contact", {})
            if any(contact.get(k) for k in ("adresse", "telephone", "email")):
                if contact_is_visible(diagnostic):
                    doc.add_paragraph(
                        f"{'Contact' if lang == 'fr' else 'Contact'}: {contact.get('adresse', '')} "
                        f"— {contact.get('telephone', '')} — {contact.get('email', '')}")
                else:
                    doc.add_paragraph(("Contact masqué" if lang == "fr" else "Contact masked"))

    _add_entreprise_section_word(doc, diagnostic, lang)

    schema = load_schema()["branches"]
    etoile = diagnostic.get("etoile", {})

    for branch_key, branch in schema.items():
        doc.add_heading(branch["label"].get(lang, branch["label"]["fr"]), level=1)
        branch_data = etoile.get(branch_key, {})
        for field in branch["fields"]:
            label = field["label"].get(lang, field["label"]["fr"])
            value = branch_data.get(field["id"], "")
            if field["type"] == "activity_list":
                if value:
                    doc.add_paragraph(label + " :", style="Intense Quote")
                    for act in value:
                        doc.add_paragraph(
                            f"{act.get('nom', '')} — part de marché relative : "
                            f"{act.get('part_marche_relative', '')}, croissance : "
                            f"{act.get('taux_croissance', '')}%",
                            style="List Bullet")
                continue
            if value in (None, "", 0, 0.0):
                continue
            unit = field.get(f"unit_{lang}") or field.get("unit_fr", "")
            value_str = f"{value} {unit}".strip() if field["type"] == "number" else str(value)
            p = doc.add_paragraph()
            p.add_run(f"{label} : ").bold = True
            p.add_run(value_str)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_diagnostic_pdf_bytes(diagnostic: dict, lang: str = "fr",
                                 include_real_name: bool = False) -> bytes:
    """Exporte les données brutes du diagnostic (étoile du conseil) en PDF."""
    from modules.collecte import load_schema

    title = "Diagnostic - Etoile du conseil" if lang == "fr" else "Diagnostic - Advisory star"
    pdf = _PDF()
    pdf.title_text = title
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    pdf.line(_identity_label(diagnostic, lang, include_real_name))
    pdf.line(f"Type: {diagnostic.get('type', '')}")
    pdf.line(f"{'Conseiller' if lang == 'fr' else 'Advisor'}: {diagnostic.get('conseiller', '')}")

    if include_real_name:
        ident = diagnostic.get("identification", {})
        if ident:
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 12)
            pdf.title_line("Identification & localisation" if lang == "fr" else "Identification & location")
            pdf.set_font("Helvetica", size=10)
            loc_fields = [
                ("Village", ident.get("village", "")),
                ("Arrondissement" if lang == "fr" else "Sub-district", ident.get("arrondissement", "")),
                ("Code arrondissement" if lang == "fr" else "Sub-district code",
                 ident.get("code_arrondissement", "")),
                ("Departement" if lang == "fr" else "Department", ident.get("departement", "")),
                ("Code departement" if lang == "fr" else "Department code",
                 ident.get("code_departement", "")),
                ("Region" if lang == "fr" else "Region", ident.get("region", "")),
                ("Code region" if lang == "fr" else "Region code", ident.get("code_region", "")),
                ("Pays" if lang == "fr" else "Country", ident.get("pays", "")),
                ("Annee" if lang == "fr" else "Year", ident.get("annee", "")),
            ]
            for label, value in loc_fields:
                if value:
                    pdf.line(f"- {label}: {value}", height=6)
            gps = ident.get("gps", {})
            if gps.get("latitude") or gps.get("longitude"):
                pdf.line(f"- GPS: {gps.get('latitude', '')}, {gps.get('longitude', '')}", height=6)

            from modules.identification import contact_is_visible
            contact = ident.get("contact", {})
            if any(contact.get(k) for k in ("adresse", "telephone", "email")):
                if contact_is_visible(diagnostic):
                    pdf.line(f"- Contact: {contact.get('adresse', '')} - "
                             f"{contact.get('telephone', '')} - {contact.get('email', '')}", height=6)
                else:
                    pdf.line("- " + ("Contact masque" if lang == "fr" else "Contact masked"), height=6)

    pdf.ln(4)

    _add_entreprise_section_pdf(pdf, diagnostic, lang)

    schema = load_schema()["branches"]
    etoile = diagnostic.get("etoile", {})

    for branch_key, branch in schema.items():
        pdf.set_font("Helvetica", "B", 13)
        pdf.title_line(branch["label"].get(lang, branch["label"]["fr"]))
        pdf.set_font("Helvetica", size=10)
        branch_data = etoile.get(branch_key, {})
        any_field = False
        for field in branch["fields"]:
            label = field["label"].get(lang, field["label"]["fr"])
            value = branch_data.get(field["id"], "")
            if field["type"] == "activity_list":
                for act in value or []:
                    any_field = True
                    pdf.line(f"- {label} : {act.get('nom', '')} (part {act.get('part_marche_relative', '')}, "
                             f"croissance {act.get('taux_croissance', '')}%)", height=6)
                continue
            if value in (None, "", 0, 0.0):
                continue
            any_field = True
            unit = field.get(f"unit_{lang}") or field.get("unit_fr", "")
            value_str = f"{value} {unit}".strip() if field["type"] == "number" else str(value)
            pdf.line(f"- {label} : {value_str}", height=6)
        if not any_field:
            pdf.line("-", height=6)
        pdf.ln(2)

    return bytes(pdf.output())
